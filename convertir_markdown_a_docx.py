#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para convertir archivos Markdown a DOCX
Convierte todos los archivos de la carpeta docs/manuales a formato .docx
"""

import os
import re
from pathlib import Path
import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

def markdown_to_docx(md_file, docx_file):
    """Convierte un archivo Markdown a DOCX"""
    
    print(f"üîÑ Convirtiendo: {os.path.basename(md_file)}")
    
    # Leer archivo Markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Crear documento Word
    doc = Document()
    
    # Configurar estilos
    styles = doc.styles
    
    # Estilo para t√≠tulos
    if 'Heading 1' not in [style.name for style in styles]:
        heading1_style = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        heading1_style.font.size = Pt(18)
        heading1_style.font.bold = True
    
    # Dividir contenido por l√≠neas
    lines = md_content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            # L√≠nea vac√≠a
            doc.add_paragraph()
            i += 1
            continue
            
        # T√≠tulos
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            title_text = line.lstrip('#').strip()
            
            if level == 1:
                heading = doc.add_heading(title_text, level=1)
            elif level == 2:
                heading = doc.add_heading(title_text, level=2)
            elif level == 3:
                heading = doc.add_heading(title_text, level=3)
            else:
                heading = doc.add_heading(title_text, level=4)
                
        # Listas
        elif line.startswith('- ') or line.startswith('* '):
            item_text = line[2:].strip()
            doc.add_paragraph(item_text, style='List Bullet')
            
        elif re.match(r'^\d+\.', line):
            item_text = re.sub(r'^\d+\.\s*', '', line)
            doc.add_paragraph(item_text, style='List Number')
            
        # C√≥digo
        elif line.startswith('```'):
            # Buscar el final del bloque de c√≥digo
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            if code_lines:
                # Crear p√°rrafo con formato de c√≥digo
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                # Fondo gris para c√≥digo
                p.paragraph_format.left_indent = Inches(0.5)
                
        # Tablas (b√°sico)
        elif '|' in line and line.count('|') > 1:
            # Procesar tabla
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells and not all(cell.startswith('-') for cell in cells):
                # Crear tabla
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for j, cell in enumerate(cells):
                    hdr_cells[j].text = cell
                    
        # Texto normal
        else:
            if line:
                # Procesar texto con formato b√°sico
                paragraph = doc.add_paragraph()
                
                # Buscar texto en negrita
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = paragraph.add_run(part[2:-2])
                        run.bold = True
                    else:
                        paragraph.add_run(part)
        
        i += 1
    
    # Guardar documento
    doc.save(docx_file)
    return True

def main():
    """Funci√≥n principal"""
    print("üöÄ Iniciando conversi√≥n de archivos Markdown a DOCX...")
    
    # Directorios
    source_dir = Path('docs/manuales')
    output_dir = Path('docs/manuales/docx')
    
    # Crear directorio de salida
    output_dir.mkdir(exist_ok=True)
    print(f"üìÅ Directorio de salida: {output_dir}")
    
    # Archivos a convertir
    files = [
        'AJUSTES_MEJORAS_SISTEMA.md',
        'CHANGELOG.md', 
        'COMENTARIOS_README_CODIGO.md',
        'DIAGRAMAS_UML_TECNICOS.md',
        'EVIDENCIA_INTEGRACION_MODULOS.md',
        'GUIA_INSTALACION_COMPLETA.md',
        'INFORME_CONTROL_AVANCES.md',
        'MANUAL_DESARROLLADOR_ARQUITECTURA.md',
        'PRUEBAS_INTERNAS.md'
    ]
    
    success_count = 0
    error_count = 0
    
    print(f"\nüìã Convirtiendo {len(files)} archivos...")
    
    for file in files:
        input_file = source_dir / file
        output_file = output_dir / file.replace('.md', '.docx')
        
        if input_file.exists():
            try:
                markdown_to_docx(input_file, output_file)
                
                if output_file.exists():
                    file_size = output_file.stat().st_size
                    file_size_kb = round(file_size / 1024, 2)
                    print(f"  ‚úÖ Convertido: {file.replace('.md', '.docx')} ({file_size_kb} KB)")
                    success_count += 1
                else:
                    print(f"  ‚ùå Error: Archivo no se cre√≥")
                    error_count += 1
                    
            except Exception as e:
                print(f"  ‚ùå Error al convertir {file}: {str(e)}")
                error_count += 1
        else:
            print(f"  ‚ö†Ô∏è  Archivo no encontrado: {input_file}")
            error_count += 1
    
    print(f"\nüìä Resumen de conversi√≥n:")
    print(f"  ‚úÖ Exitosos: {success_count}")
    print(f"  ‚ùå Errores: {error_count}")
    print(f"  üìÅ Archivos guardados en: {output_dir}")
    
    if success_count > 0:
        print(f"\nüéâ ¬°Conversi√≥n completada!")
        print(f"üí° Los archivos .docx est√°n listos para usar en Microsoft Word")
        
        # Abrir carpeta de destino
        try:
            import subprocess
            subprocess.run(['explorer', str(output_dir)], check=True)
            print(f"üìÇ Carpeta de destino abierta")
        except:
            print(f"üìÇ Carpeta de destino: {output_dir}")
    else:
        print(f"\n‚ö†Ô∏è  No se pudieron convertir archivos")

if __name__ == "__main__":
    main()









